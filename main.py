from fastapi import FastAPI
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from docx import Document
from io import BytesIO
from pathlib import Path
from datetime import datetime, timedelta
import re
import os
import json
import base64
import urllib.request
import urllib.error
import time
import uuid

app = FastAPI(
    title="Entry Pass API",
    version="1.0.0",
    description="خدمة مستقلة لإنشاء استمارة سمة الدخول بصيغة Word"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "entry_pass_template.docx"
HOTEL_TEMPLATE_PATH = BASE_DIR / "hotel_booking_template.docx"
GEMMA_MODEL = "gemma-4-26b-a4b-it"
GEMMA_API_KEY = os.getenv("GEMMA_API_KEY", "").strip()
SUPABASE_URL = os.getenv("SUPABASE_URL", "https://hifkuvyvhrxmcgkbvgqo.supabase.co").rstrip("/")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "").strip()
GEMMA_INPUT_COST_PER_MILLION = float(os.getenv("GEMMA_INPUT_COST_PER_MILLION", "0") or 0)
GEMMA_OUTPUT_COST_PER_MILLION = float(os.getenv("GEMMA_OUTPUT_COST_PER_MILLION", "0") or 0)



class EntryPassRequest(BaseModel):
    nationality: str = Field(default="الباكستانية")
    count: int = Field(ge=1, le=500)
    first_name: str
    last_name: str
    entry_port: str = ""
    hotel: str = ""
    arrival_date: str = ""
    marketing_company: str = ""
    telegram_destination: str = ""
    document_date: str = ""
    document_number: int = Field(ge=1, le=999999999)


class HotelBookingRequest(BaseModel):
    company: str = ""
    first_name: str = ""
    last_name: str = ""
    count: int = Field(ge=1, le=500)
    hotel: str = ""
    document_date: str = ""
    arrival_date: str = ""

class GemmaReadRequest(BaseModel):
    image_base64: str
    mime_type: str = "image/jpeg"



def _clean(value: str) -> str:
    value = (value or "").strip()
    value = re.sub(r"\s+", " ", value)
    return value


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    # القالب الذي نولده يحفظ كل placeholder ضمن Run واحد.
    # هذا المسار أيضاً يعالج الحالات التي ينقسم بها النص بين عدة Runs.
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = full_text
    changed = False
    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, value)
            changed = True

    if not changed:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _replace_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replacements)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, replacements)
        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, replacements)


@app.get("/")
def root():
    return {
        "ok": True,
        "service": "entry-pass-api",
        "message": "سيرفر سمة الدخول يعمل"
    }


@app.get("/health")
def health():
    return {
        "ok": TEMPLATE_PATH.exists() and HOTEL_TEMPLATE_PATH.exists(),
        "entry_template": TEMPLATE_PATH.name,
        "entry_template_exists": TEMPLATE_PATH.exists(),
        "hotel_template": HOTEL_TEMPLATE_PATH.name,
        "hotel_template_exists": HOTEL_TEMPLATE_PATH.exists(),
        "gemma_model": GEMMA_MODEL,
        "gemma_configured": bool(GEMMA_API_KEY or SUPABASE_SERVICE_ROLE_KEY),
    }


@app.post("/fill-entry-pass")
def fill_entry_pass(data: EntryPassRequest):
    if not TEMPLATE_PATH.exists():
        return {
            "ok": False,
            "error": "قالب Word غير موجود على السيرفر"
        }

    tomorrow = datetime.now() + timedelta(days=1)
    arrival_default = tomorrow + timedelta(days=15)

    document_date = _clean(data.document_date)
    if not document_date:
        document_date = tomorrow.strftime("%d / %m / %Y")

    arrival_date = _clean(data.arrival_date)
    if not arrival_date:
        arrival_date = arrival_default.strftime("%d / %m / %Y")

    replacements = {
        "{{DOCUMENT_NUMBER}}": str(data.document_number),
        "{{NUM}}": str(data.document_number),
        "{{NATIONALITY}}": _clean(data.nationality),
        "{{NAT}}": _clean(data.nationality),
        "{{COUNT}}": str(data.count),
        "{{CNT}}": str(data.count),
        "{{FIRST_NAME}}": _clean(data.first_name),
        "{{FIRST}}": _clean(data.first_name),
        "{{LAST_NAME}}": _clean(data.last_name),
        "{{LAST}}": _clean(data.last_name),
        "{{ENTRY_PORT}}": _clean(data.entry_port),
        "{{PORT}}": _clean(data.entry_port),
        "{{HOTEL}}": _clean(data.hotel),
        "{{ARRIVAL_DATE}}": arrival_date,
        "{{ARR}}": arrival_date,
        "{{MARKETING_COMPANY}}": _clean(data.marketing_company),
        "{{COMP}}": _clean(data.marketing_company),
        "{{TELEGRAM_DESTINATION}}": _clean(data.telegram_destination),
        "{{TEL}}": _clean(data.telegram_destination),
        "{{DOCUMENT_DATE}}": document_date,
        "{{DATE}}": document_date,
    }

    doc = Document(str(TEMPLATE_PATH))
    _replace_everywhere(doc, replacements)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"entry-pass-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"

    return StreamingResponse(
        output,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Entry-Pass-Count": str(data.count),
            "X-Entry-Pass-Document-Number": str(data.document_number),
        },
    )


@app.post("/fill-hotel-booking")
def fill_hotel_booking(data: HotelBookingRequest):
    if not HOTEL_TEMPLATE_PATH.exists():
        return {
            "ok": False,
            "error": "قالب الحجز الفندقي غير موجود على السيرفر"
        }

    tomorrow = datetime.now() + timedelta(days=1)
    arrival_default = tomorrow + timedelta(days=15)
    document_date = _clean(data.document_date) or tomorrow.strftime("%d / %m / %Y")
    arrival_date = _clean(data.arrival_date) or arrival_default.strftime("%d / %m / %Y")

    replacements = {
        "{{COMPANY}}": _clean(data.company),
        "{{FIRST_NAME}}": _clean(data.first_name),
        "{{LAST_NAME}}": _clean(data.last_name),
        "{{COUNT}}": str(data.count),
        "{{HOTEL}}": _clean(data.hotel),
        "{{DOCUMENT_DATE}}": document_date,
        "{{ARRIVAL_DATE}}": arrival_date,
    }

    doc = Document(str(HOTEL_TEMPLATE_PATH))
    _replace_everywhere(doc, replacements)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"hotel-booking-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
    return StreamingResponse(
        output,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Hotel-Booking-Count": str(data.count),
        },
    )


def _log_ai_usage(*, provider: str, model: str, request_type: str = "passport_read",
                  input_tokens: int = 0, output_tokens: int = 0, total_tokens: int = 0,
                  estimated_cost_usd: float = 0.0, success: bool = True,
                  error_code: str | None = None, latency_ms: int | None = None,
                  key_id: str | None = None) -> None:
    """Best-effort analytics logging. Never breaks the main passport request."""
    if not SUPABASE_SERVICE_ROLE_KEY:
        return
    row = {
        "request_id": str(uuid.uuid4()),
        "provider": provider,
        "model": model,
        "request_type": request_type,
        "input_tokens": int(input_tokens or 0),
        "output_tokens": int(output_tokens or 0),
        "total_tokens": int(total_tokens or 0),
        "estimated_cost_usd": float(estimated_cost_usd or 0),
        "success": bool(success),
        "error_code": error_code,
        "latency_ms": int(latency_ms) if latency_ms is not None else None,
        "key_id": str(key_id) if key_id else None,
    }
    try:
        req = urllib.request.Request(
            f"{SUPABASE_URL}/rest/v1/ai_usage_logs",
            data=json.dumps(row).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "apikey": SUPABASE_SERVICE_ROLE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
                "Prefer": "return=minimal",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=10):
            pass
    except Exception as exc:
        print(f"AI analytics log skipped: {exc}")


_GEMMA_KEYS_CACHE: list[dict] = []
_GEMMA_KEYS_CACHE_AT: float = 0.0
_GEMMA_KEYS_CACHE_SECONDS = 30.0


def _supabase_rest_request(path: str, *, method: str = "GET", body: dict | None = None,
                           prefer: str | None = None) -> bytes:
    if not SUPABASE_SERVICE_ROLE_KEY:
        raise RuntimeError("SUPABASE_SERVICE_ROLE_KEY is not configured")
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
    }
    payload = None
    if body is not None:
        headers["Content-Type"] = "application/json"
        payload = json.dumps(body).encode("utf-8")
    if prefer:
        headers["Prefer"] = prefer
    req = urllib.request.Request(
        f"{SUPABASE_URL}/rest/v1/{path}",
        data=payload,
        headers=headers,
        method=method,
    )
    with urllib.request.urlopen(req, timeout=12) as response:
        return response.read()


def _get_active_gemma_keys(force_refresh: bool = False) -> list[dict]:
    """Return active Gemma keys from Supabase; Render env key remains fallback."""
    global _GEMMA_KEYS_CACHE, _GEMMA_KEYS_CACHE_AT
    now = time.monotonic()
    if not force_refresh and _GEMMA_KEYS_CACHE and (now - _GEMMA_KEYS_CACHE_AT) < _GEMMA_KEYS_CACHE_SECONDS:
        return list(_GEMMA_KEYS_CACHE)

    keys: list[dict] = []
    if SUPABASE_SERVICE_ROLE_KEY:
        try:
            raw = _supabase_rest_request(
                "gemma_api_keys?select=id,api_key,label,is_active,created_at&is_active=eq.true&order=created_at.asc"
            )
            rows = json.loads(raw.decode("utf-8") or "[]")
            if isinstance(rows, list):
                for row in rows:
                    api_key = str((row or {}).get("api_key") or "").strip()
                    if api_key:
                        keys.append({
                            "id": str((row or {}).get("id") or ""),
                            "api_key": api_key,
                            "label": str((row or {}).get("label") or "Gemma Key"),
                        })
        except Exception as exc:
            print(f"Gemma key pool unavailable; using Render fallback if present: {exc}")

    # Keep the Render secret as a safe fallback without exposing it in the dashboard.
    if not keys and GEMMA_API_KEY:
        keys.append({"id": "render-env", "api_key": GEMMA_API_KEY, "label": "Render fallback"})

    _GEMMA_KEYS_CACHE = list(keys)
    _GEMMA_KEYS_CACHE_AT = now
    return keys


def _mark_gemma_key_used(key_id: str | None) -> None:
    if not key_id or key_id == "render-env" or not SUPABASE_SERVICE_ROLE_KEY:
        return
    try:
        _supabase_rest_request(
            f"gemma_api_keys?id=eq.{key_id}",
            method="PATCH",
            body={"last_used_at": datetime.utcnow().isoformat() + "Z"},
            prefer="return=minimal",
        )
    except Exception as exc:
        print(f"Could not update Gemma last_used_at: {exc}")


GEMMA_PROMPT = r"""
You are a passport data extraction engine. Analyze ONLY the passport identity page image.
Return ONLY one valid JSON object, with no markdown and no explanation.

Required keys:
{
  "given_name_en": null,
  "father_name_en": null,
  "surname_en": null,
  "passport_number": null,
  "nationality": null,
  "residence_country": null,
  "birth_date": null,
  "issue_date": null,
  "expiry_date": null,
  "sex": null,
  "mrz_line1": null,
  "mrz_line2": null,
  "confidence": 0.0
}

Rules:
- Never guess unreadable values; use null.
- Dates must be YYYY-MM-DD.
- nationality should be the 3-letter ICAO code when available, e.g. IRQ, PAK, IND.
- sex must be M, F, or X when readable.
- Copy MRZ lines exactly as visible, using < fillers.
- Compare printed fields with MRZ before returning the result.
- Prefer checksum-consistent MRZ for passport number, nationality, birth date, expiry date and sex.
- For Pakistani passports, Father/Husband Name may be printed separately. Put that value in father_name_en.
- Preserve the holder's printed English names; do not translate names here.
- confidence must be a number from 0.0 to 1.0 reflecting overall readability and agreement.
"""


def _extract_json_object(text: str) -> dict:
    raw = (text or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.I)
        raw = re.sub(r"\s*```$", "", raw)
    try:
        obj = json.loads(raw)
        if isinstance(obj, dict):
            return obj
    except Exception:
        pass

    start = raw.find("{")
    end = raw.rfind("}")
    if start >= 0 and end > start:
        obj = json.loads(raw[start:end + 1])
        if isinstance(obj, dict):
            return obj
    raise ValueError("Gemma returned invalid JSON")


@app.post("/read-gemma")
def read_gemma(data: GemmaReadRequest):
    try:
        image_bytes = base64.b64decode(data.image_base64, validate=True)
    except Exception:
        return {"ok": False, "error": "Invalid base64 image"}

    if not image_bytes:
        return {"ok": False, "error": "Empty image"}
    if len(image_bytes) > 15 * 1024 * 1024:
        return {"ok": False, "error": "Image is too large"}

    mime_type = (data.mime_type or "image/jpeg").strip().lower()
    if mime_type not in {"image/jpeg", "image/jpg", "image/png", "image/webp"}:
        mime_type = "image/jpeg"

    payload = {
        "contents": [{
            "role": "user",
            "parts": [
                {"text": GEMMA_PROMPT},
                {
                    "inline_data": {
                        "mime_type": mime_type,
                        "data": base64.b64encode(image_bytes).decode("ascii")
                    }
                }
            ]
        }],
        "generationConfig": {
            "temperature": 0.0,
            "thinkingConfig": {"thinkingLevel": "minimal"}
        }
    }

    keys = _get_active_gemma_keys()
    if not keys:
        return {
            "ok": False,
            "error": "No active Gemma key. Add one from the Admin Dashboard or configure GEMMA_API_KEY on Render."
        }

    url = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        f"{GEMMA_MODEL}:generateContent"
    )
    errors: list[str] = []

    for key_ref in keys:
        started_at = time.perf_counter()
        key_id = str(key_ref.get("id") or "")
        key_label = str(key_ref.get("label") or "Gemma Key")
        api_key = str(key_ref.get("api_key") or "").strip()
        if not api_key:
            continue

        request = urllib.request.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "x-goog-api-key": api_key,
            },
            method="POST",
        )

        try:
            with urllib.request.urlopen(request, timeout=90) as response:
                raw_response = response.read().decode("utf-8")
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")
            latency_ms = int((time.perf_counter() - started_at) * 1000)
            _log_ai_usage(
                provider="gemma", model=GEMMA_MODEL, success=False,
                error_code=str(exc.code), latency_ms=latency_ms, key_id=key_id,
            )
            errors.append(f"{key_label}: HTTP {exc.code} - {detail[:250]}")
            continue
        except Exception as exc:
            latency_ms = int((time.perf_counter() - started_at) * 1000)
            _log_ai_usage(
                provider="gemma", model=GEMMA_MODEL, success=False,
                error_code="connection_error", latency_ms=latency_ms, key_id=key_id,
            )
            errors.append(f"{key_label}: connection error - {exc}")
            continue

        try:
            api_json = json.loads(raw_response)
            parts = (
                api_json.get("candidates", [{}])[0]
                .get("content", {})
                .get("parts", [])
            )
            output_text = "\n".join(
                str(part.get("text", ""))
                for part in parts
                if isinstance(part, dict) and part.get("text")
            ).strip()
            result = _extract_json_object(output_text)
        except Exception as exc:
            latency_ms = int((time.perf_counter() - started_at) * 1000)
            _log_ai_usage(
                provider="gemma", model=GEMMA_MODEL, success=False,
                error_code="parse_error", latency_ms=latency_ms, key_id=key_id,
            )
            errors.append(f"{key_label}: parse error - {exc}")
            continue

        usage = api_json.get("usageMetadata", {}) or {}
        input_tokens = int(usage.get("promptTokenCount") or 0)
        output_tokens = int(usage.get("candidatesTokenCount") or 0)
        total_tokens = int(usage.get("totalTokenCount") or (input_tokens + output_tokens))
        estimated_cost = (
            (input_tokens / 1_000_000) * GEMMA_INPUT_COST_PER_MILLION
            + (output_tokens / 1_000_000) * GEMMA_OUTPUT_COST_PER_MILLION
        )
        latency_ms = int((time.perf_counter() - started_at) * 1000)

        _log_ai_usage(
            provider="gemma", model=GEMMA_MODEL,
            input_tokens=input_tokens, output_tokens=output_tokens,
            total_tokens=total_tokens, estimated_cost_usd=estimated_cost,
            success=True, latency_ms=latency_ms, key_id=key_id,
        )
        _mark_gemma_key_used(key_id)

        return {
            "ok": True,
            "model": GEMMA_MODEL,
            "data": result,
            "usage": {
                "key_id": key_id,
                "key_label": key_label,
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "total_tokens": total_tokens,
                "estimated_cost_usd": estimated_cost,
                "latency_ms": latency_ms,
            },
        }

    return {
        "ok": False,
        "error": "All active Gemma keys failed",
        "detail": "\n---\n".join(errors)[-5000:],
    }

